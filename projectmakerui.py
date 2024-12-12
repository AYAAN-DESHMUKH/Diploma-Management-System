import tkinter as tk
from tkinter import scrolledtext
from docx.shared import Inches, Pt
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from tkinter import messagebox



class ProjectInfoUI:
    
    def __init__(self, master):
        self.master = master
        self.master.title("Project Information Form")
        self.master.configure(bg="black")
        self.master.attributes("-fullscreen", True)
        self.master.bind("<Escape>", lambda event: self.master.attributes("-fullscreen", False))

        self.create_widgets()

    def create_widgets(self):
        # Topic
        tk.Label(self.master, text="Topic:", bg="black", fg="orange", font=("Arial", 14)).grid(row=0, column=0, padx=10, pady=5)
        self.topic_entry = tk.Entry(self.master, width=50, bg="black", fg="orange", font=("Arial", 14))
        self.topic_entry.grid(row=0, column=1, columnspan=2, pady=5)

        # Guidance Name
        tk.Label(self.master, text="Guidance Name:", bg="black", fg="orange", font=("Arial", 14)).grid(row=1, column=0, padx=10, pady=5)
        self.guidance_entry = tk.Entry(self.master, width=50, bg="black", fg="orange", font=("Arial", 14))
        self.guidance_entry.grid(row=1, column=1, columnspan=2, pady=5)

        # Group Members
        tk.Label(self.master, text="Group Members:", bg="black", fg="orange", font=("Arial", 14)).grid(row=2, column=0, padx=10, pady=5)

        self.member1_entry = tk.Entry(self.master, width=15, bg="black", fg="orange", font=("Arial", 14))
        self.member1_entry.grid(row=2, column=1, pady=5, padx=(10, 5))

        self.member2_entry = tk.Entry(self.master, width=15, bg="black", fg="orange", font=("Arial", 14))
        self.member2_entry.grid(row=2, column=2, pady=5, padx=5)

        self.member3_entry = tk.Entry(self.master, width=15, bg="black", fg="orange", font=("Arial", 14))
        self.member3_entry.grid(row=2, column=3, pady=5, padx=5)

        self.member4_entry = tk.Entry(self.master, width=15, bg="black", fg="orange", font=("Arial", 14))
        self.member4_entry.grid(row=2, column=4, pady=5, padx=(5, 10))

        # Abstract
        tk.Label(self.master, text="Abstract:", bg="black", fg="orange", font=("Arial", 14)).grid(row=3, column=0, padx=10, pady=5)
        self.abstract_text = scrolledtext.ScrolledText(self.master, width=50, height=5, bg="black", fg="orange", font=("Arial", 14))
        self.abstract_text.grid(row=3, column=1, columnspan=4, pady=5)

        # Introduction
        tk.Label(self.master, text="Introduction:", bg="black", fg="orange", font=("Arial", 14)).grid(row=4, column=0, padx=10, pady=5)
        self.intro_text = scrolledtext.ScrolledText(self.master, width=50, height=5, bg="black", fg="orange", font=("Arial", 14))
        self.intro_text.grid(row=4, column=1, columnspan=4, pady=5)

        # Conclusion
        tk.Label(self.master, text="Conclusion:", bg="black", fg="orange", font=("Arial", 14)).grid(row=5, column=0, padx=10, pady=5)
        self.conclusion_text = scrolledtext.ScrolledText(self.master, width=50, height=5, bg="black", fg="orange", font=("Arial", 14))
        self.conclusion_text.grid(row=5, column=1, columnspan=4, pady=5)

        # Reference
        tk.Label(self.master, text="Reference:", bg="black", fg="orange", font=("Arial", 14)).grid(row=6, column=0, padx=10, pady=5)
        self.reference_text = scrolledtext.ScrolledText(self.master, width=50, height=5, bg="black", fg="orange", font=("Arial", 14))
        self.reference_text.grid(row=6, column=1, columnspan=4, pady=5)

        # Submit Button
        submit_button = tk.Button(self.master, text="Submit", command=self.submit_form, bg="orange", fg="black", font=("Arial", 14))
        submit_button.grid(row=7, column=2, pady=10)
        submit_button1 = tk.Button(self.master, text="Back", command=self.home, bg="orange", fg="black", font=("Arial", 14))
        submit_button1.grid(row=8, column=2, padx=50,pady=10)
    def submit_form(self):
        
        # Retrieve data from input fields
        topic = self.topic_entry.get()
        guidance_name = self.guidance_entry.get()
        member1 = self.member1_entry.get()
        member2 = self.member2_entry.get()
        member3 = self.member3_entry.get()
        member4 = self.member4_entry.get()
        abstract = self.abstract_text.get("1.0", tk.END)
        introduction = self.intro_text.get("1.0", tk.END)
        conclusion = self.conclusion_text.get("1.0", tk.END)
        reference = self.reference_text.get("1.0", tk.END)

        # Print or process the retrieved data (you can modify this part based on your requirements)
        print("Topic:", topic)
        print("Guidance Name:", guidance_name)
        print("Group Members:")
        print("   Member 1:", member1)
        print("   Member 2:", member2)
        print("   Member 3:", member3)
        print("   Member 4:", member4)
        print("Abstract:", abstract)
        print("Introduction:", introduction)
        print("Conclusion:", conclusion)
        print("Reference:", reference)
        
        doc = Document()
    
        # Get the section properties element
        sec_pr = doc.sections[0]._sectPr if doc.sections else doc.sections.add()._sectPr

        # Create new page borders element
        pg_borders = OxmlElement('w:pgBorders')
        pg_borders.set(qn('w:offsetFrom'), 'page')

        # Set individual borders
        for border_name in ('top', 'left', 'bottom', 'right'):
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')  # a single line
            border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
            border_el.set(qn('w:space'), '24')  # space
            border_el.set(qn('w:color'), 'auto')  # color
            pg_borders.append(border_el)
        
        # Append page borders to section properties
        sec_pr.append(pg_borders)

        # Set the document alignment to center
        for section in doc.sections:
            for paragraph in section.footer.paragraphs:
                for run in paragraph.runs:
                    run.paragraph.alignment = 1

        # Adjust the left margin to center the image
        for section in doc.sections:
            section.left_margin = Inches((8.5 - 7.0) / 2)  # Assuming default page width is 8.5 inches

        # Add a new paragraph with centered alignment
        paragraph = doc.add_paragraph('')
        paragraph.alignment = 1

        # Add the image to the paragraph
        image_path = "C:/Users/Ayaan Deshmukh/Downloads/kalsekar.png"  # Replace with the actual path to your image
        paragraph.add_run().add_picture(image_path, width=Inches(7.0))

        # Add a new paragraph with project details
        project_title_paragraph = doc.add_paragraph("")
        project_title_paragraph.add_run("A Major-Project Proposal\nOn\n").bold = False
        project_title_paragraph.runs[0].font.italic = True
        project_title_paragraph.paragraph_format.alignment = 1

        # Add the run and set its font size to 22pt
        run_project_title = project_title_paragraph.add_run(topic)
        run_project_title.bold = True
        run_project_title.font.size = Pt(22)

        # Add a new paragraph with team members
        team_members_paragraph = doc.add_paragraph("")
        team_members_paragraph.add_run("By\n").bold = True
        team_members_paragraph.add_run("1."+member1+"\n").bold = False
        team_members_paragraph.add_run("2."+member2+"\n").bold = False
        team_members_paragraph.add_run("3."+member3+"\n").bold = False
        team_members_paragraph.add_run("4."+member4+"\n").bold = False
        team_members_paragraph.paragraph_format.alignment = 1  # Align center
        team_members_paragraph.paragraph_format.space_after = Pt(18)  # Add spacing

        # Add a new paragraph with guidance
        guidance_paragraph = doc.add_paragraph("")
        guidance_paragraph.add_run("Under Guidance of\n").bold = True
        guidance_paragraph.add_run(guidance_name).bold = False
        guidance_paragraph.paragraph_format.alignment = 1  # Align center
        guidance_paragraph.paragraph_format.space_after = Pt(18)  # Add spacing

        # Add a new paragraph with course details
        course_details_paragraph = doc.add_paragraph("")
        course_details_paragraph.add_run("In\n").bold = True
        course_details_paragraph.add_run("Three Years Diploma Program in Engineering & Technology of\n").bold = False
        course_details_paragraph.add_run("Maharashtra State Board of Technical Education, Mumbai (Autonomous)\n").bold = False
        course_details_paragraph.add_run("ISO 9001:2015\n").bold = False
        course_details_paragraph.paragraph_format.alignment = 1  # Align center
        course_details_paragraph.paragraph_format.space_after = Pt(18)  # Add spacing

        # Set font size for the entire document
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(19)

        """
        
        second page
        
        """

        doc.add_page_break()
        # Get the section properties element

        # Create new page borders element
        pg_borders = OxmlElement('w:pgBorders')
        pg_borders.set(qn('w:offsetFrom'), 'page')

        # Set individual borders
        for border_name in ('top', 'left', 'bottom', 'right'):
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')  # a single line
            border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
            border_el.set(qn('w:space'), '24')  # space
            border_el.set(qn('w:color'), 'auto')  # color
            pg_borders.append(border_el)
        # Append page borders to section properties
        sec_pr.append(pg_borders)
        paragraph2 = doc.add_paragraph('')
        paragraph2.alignment = 1
        image_pat = "C:/Users/Ayaan Deshmukh/Downloads/cert.png"  # Replace with the actual path to your image
        paragraph2.add_run().add_picture(image_pat, width=Inches(7.0))

        """
        third page
        """

        # Create new page borders element
        pg_borders = OxmlElement('w:pgBorders')
        pg_borders.set(qn('w:offsetFrom'), 'page')

        # Set individual borders
        for border_name in ('top', 'left', 'bottom', 'right'):
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')  # a single line
            border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
            border_el.set(qn('w:space'), '24')  # space
            border_el.set(qn('w:color'), 'auto')  # color
            pg_borders.append(border_el)
        # Append page borders to section properties
        sec_pr.append(pg_borders)
        paragraph3 = doc.add_paragraph('')
        paragraph3.alignment = 1
        image_pat3 = "C:/Users/Ayaan Deshmukh/Downloads/annexure1.png"  # Replace with the actual path to your image
        paragraph3.add_run().add_picture(image_pat3, width=Inches(7.0))

        """
        fourth page
        """

        # Create new page borders element
        paragraph3.alignment = 1
        image_pat3 = "C:/Users/Ayaan Deshmukh/Downloads/ann2.png"  # Replace with the actual path to your image
        paragraph3.add_run().add_picture(image_pat3, width=Inches(7.0))

        """
        fifth page
        """

        # Create new page borders element
        pg_borders = OxmlElement('w:pgBorders')
        pg_borders.set(qn('w:offsetFrom'), 'page')

        # Set individual borders
        for border_name in ('top', 'left', 'bottom', 'right'):
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')  # a single line
            border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
            border_el.set(qn('w:space'), '24')  # space
            border_el.set(qn('w:color'), 'auto')  # color
            pg_borders.append(border_el)
        # Append page borders to section properties
        sec_pr.append(pg_borders)
        paragraph3 = doc.add_paragraph('')
        paragraph3.alignment = 1
        image_pat3 = "C:/Users/Ayaan Deshmukh/Downloads/ann3.png"  # Replace with the actual path to your image
        paragraph3.add_run().add_picture(image_pat3, width=Inches(7.0))

        # new page

        # Add a new section for the next page
        doc.add_section()

        # Get the section properties element for the new section
        new_sec_pr = doc.sections[-1]._sectPr

        # Create new page borders element for the new section
        new_pg_borders = OxmlElement('w:pgBorders')
        new_pg_borders.set(qn('w:offsetFrom'), 'page')

        # Set individual borders for the new section
        for border_name in ('top', 'left', 'bottom', 'right'):
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')  # a single line
            border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
            border_el.set(qn('w:space'), '24')  # space
            border_el.set(qn('w:color'), 'auto')  # color
            new_pg_borders.append(border_el)

        # Append new page borders to section properties for the new section
        new_sec_pr.append(new_pg_borders)

        # Add a new centered paragraph with the title "Introduction"
        introduction_paragraph = doc.add_paragraph("Introduction")
        introduction_paragraph.paragraph_format.alignment = 1  # Center alignment
        introduction_paragraph.runs[0].font.size = Pt(32)  # Set font size to 32

        # Add a new paragraph below the introduction with font size 12
        text_below_paragraph = doc.add_paragraph(introduction)
        text_below_paragraph.runs[0].font.size = Pt(12)  # Set font size to 12
        paragraph_format = text_below_paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(23)
        doc.add_page_break()

        # new page abstract page
        # Get the section properties element for the new section
        new_sec_pr = doc.sections[-1]._sectPr

        # Create new page borders element for the new section
        new_pg_borders = OxmlElement('w:pgBorders')
        new_pg_borders.set(qn('w:offsetFrom'), 'page')

        # Set individual borders for the new section
        for border_name in ('top', 'left', 'bottom', 'right'):
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')  # a single line
            border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
            border_el.set(qn('w:space'), '24')  # space
            border_el.set(qn('w:color'), 'auto')  # color
            new_pg_borders.append(border_el)

        # Append new page borders to section properties for the new section
        new_sec_pr.append(new_pg_borders)

        # Add a new centered paragraph with the title "Introduction"
        introduction_paragraph = doc.add_paragraph("Abstract")
        introduction_paragraph.paragraph_format.alignment = 1  # Center alignment
        introduction_paragraph.runs[0].font.size = Pt(32)  # Set font size to 32

        # Add a new paragraph below the introduction with font size 12
        text_below_paragraph = doc.add_paragraph(abstract)
        text_below_paragraph.runs[0].font.size = Pt(12)  # Set font size to 12
        paragraph_format = text_below_paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(23)

        doc.add_page_break()

        new_sec_pr = doc.sections[-1]._sectPr

        # Create new page borders element for the new section
        new_pg_borders = OxmlElement('w:pgBorders')
        new_pg_borders.set(qn('w:offsetFrom'), 'page')

        # Set individual borders for the new section
        for border_name in ('top', 'left', 'bottom', 'right'):

            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')  # a single line
            border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
            border_el.set(qn('w:space'), '24')  # space
            border_el.set(qn('w:color'), 'auto')  # color
            new_pg_borders.append(border_el)

        # Append new page borders to section properties for the new section
        new_sec_pr.append(new_pg_borders)

        # Add a new centered paragraph with the title "Introduction"
        introduction_paragraph = doc.add_paragraph("Reference")
        introduction_paragraph.paragraph_format.alignment = 1  # Center alignment
        introduction_paragraph.runs[0].font.size = Pt(32)  # Set font size to 32

        # Add a new paragraph below the introduction with font size 12
        text_below_paragraph = doc.add_paragraph(reference)
        text_below_paragraph.runs[0].font.size = Pt(12)  # Set font size to 12
        paragraph_format = text_below_paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(23)
        
        
        doc.add_page_break()

        new_sec_pr = doc.sections[-1]._sectPr

        # Create new page borders element for the new section
        new_pg_borders = OxmlElement('w:pgBorders')
        new_pg_borders.set(qn('w:offsetFrom'), 'page')

        # Set individual borders for the new section
        for border_name in ('top', 'left', 'bottom', 'right'):

            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')  # a single line
            border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
            border_el.set(qn('w:space'), '24')  # space
            border_el.set(qn('w:color'), 'auto')  # color
            new_pg_borders.append(border_el)
 
        # Append new page borders to section properties for the new section
        new_sec_pr.append(new_pg_borders)

        # Add a new centered paragraph with the title "Introduction"
        introduction_paragraph = doc.add_paragraph("Conclusion")
        introduction_paragraph.paragraph_format.alignment = 1  # Center alignment
        introduction_paragraph.runs[0].font.size = Pt(32)  # Set font size to 32

        # Add a new paragraph below the introduction with font size 12
        text_below_paragraph = doc.add_paragraph(conclusion)
        text_below_paragraph.runs[0].font.size = Pt(12)  # Set font size to 12
        paragraph_format = text_below_paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(23)
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        doc_path = os.path.join(desktop_path, topic+".docx")
        doc.save(doc_path)
        
        #successfull message
        messagebox.showinfo("Mission Accomplished!", "Great news! Your project named "+topic+" file has been successfully created. Now, go check your desktop to find it")
        



    def home(self):
        self.master.destroy()
        
    

if __name__ == "__main__":
    root = tk.Tk()
    project_info_ui = ProjectInfoUI(root)
    root.mainloop()

