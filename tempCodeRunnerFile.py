import tkinter as tk
from tkinter import filedialog, ttk
from PyPDF2 import PdfReader
import spacy
from spacy.lang.en import English
from docx import Document
import os
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from tkinter import messagebox
import webbrowser

class QB:
    def explore_modal_paper(self):
        # Action for Modal Answer Paper button
        messagebox.showinfo("Explore Modal Answer Paper", "Redirecting to Official MSBTE Modal Answer Papers website.")
        webbrowser.open("https://msbte.org.in/portal/model-answer-search/")
        
    def explore_question_papers(self):
        # Action for Question Papers button
        messagebox.showinfo("Explore Question Papers", "Redirecting to Official MSBTE Exam Question Papers website.")
        webbrowser.open("https://msbte.org.in/portal/model-answer-search/")
   
    def __init__(self, master):
           
        
           
        self.master = master
        master.title("Question Maker")
        master.configure(bg='black')
        master.attributes('-fullscreen', True)  # Full-screen mode

        # Create and place the title label
        self.title_label = ttk.Label(master, text="Question Maker", font=('Arial', 34), background='black',
                                     foreground='white')
        self.title_label.place(x=50, y=70)

        # Create and place the content
        self.content = ttk.Label(master, text=""" Question Maker simplifies the creation of a question bank from a PDF. \n Select your PDF, wait a moment, and voila! Explore generated questions in a Word document on your desktop. Happy learning!""", font=('Arial', 14), background='black',
                                     foreground='white')
        self.content.place(x=80, y=140)

        # Create and place a button to select a PDF file
        select_button = tk.Button(master, text="Select PDF", command=self.on_select_pdf, bg='black', fg='white',
                                  font=('Arial', 24))
        select_button.place(relx=0.5, rely=0.5, anchor=tk.CENTER)

        # Create and place buttons for Modal Answer Paper and Question Papers
        modal_button = tk.Button(master, text="Modal Answer Paper", command=self.explore_modal_paper, bg='black' ,fg='white',
                                 font=('Arial', 18))
        modal_button.place(relx=0.2, rely=0.5, anchor=tk.CENTER)

        question_button = tk.Button(master, text="Question Papers", command=self.explore_question_papers, bg='black', fg='white',
                                    font=('Arial', 18))
        question_button.place(relx=0.8, rely=0.5, anchor=tk.CENTER)

        # Create and place an exit button at the bottom right
        exit_button = tk.Button(master, text="Back", command=self.exit_application, bg='black', fg='white',
                                font=('Arial', 24))
        exit_button.place(x=1240, y=670)

        # Create and place a help button at the bottom
        help_button = tk.Button(master, text="Help", command=self.show_help, bg='black', fg='white',
                                font=('Arial', 24))
        help_button.place(x=70, y=670)

        # Create a Toplevel for the file dialog
        self.file_dialog_top = tk.Toplevel(self.master)
        self.file_dialog_top.withdraw()  # Hide the Toplevel
        md = ttk.Label(self.master, text="Explore Official MSBTE Modal Answer Papers", font=('Arial', 14), background='black',
                                foreground='white')
        md.place(x=85,y=420)
        
        ques= ttk.Label(self.master, text="Explore Official MSBTE Modal Answer Papers", font=('Arial', 14), background='black',
                                foreground='white')
        ques.place(x=895,y=420)

    def on_select_pdf(self):
        # Display "Just wait for a moment..." text
        loading_text = ttk.Label(self.master, text="Just wait for a moment...", font=('Arial', 18), background='black',
                                foreground='white')
        loading_text.place(relx=0.5, rely=0.6, anchor=tk.CENTER)
        self.master.update_idletasks()

        # File dialog to select a PDF file
        file_path = filedialog.askopenfilename(parent=self.file_dialog_top, filetypes=[("PDF files", "*.pdf")])

        # Destroy the Toplevel after the file is selected
        self.file_dialog_top.destroy()

        if file_path:
            # Update the text message after selecting the PDF
            loading_text.config(text="Just wait for a moment...")

            pdf_text = self.display_pdf_content(file_path)
            print("\nPDF Content:")
            print(pdf_text)

            # Generate questions based on the PDF text
            generated_questions = self.generate_questions(pdf_text)
            messagebox.showinfo("Ready for the test?",
                                "Your Question Bank has been successfully crafted. Head to your desktop to explore this treasure trove of knowledge. Happy learning!")

            # Display generated questions
            print("\nGenerated Questions:")
            for i, question in enumerate(generated_questions, start=1):
                print(f"{i}. {question}")

            # Create a Word document and save it on the desktop
            self.create_word_document(generated_questions)

    def display_pdf_content(self, file_path):
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            text = ""

            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()

            return text

    def exit_application(self):
        self.master.destroy()

    def generate_questions(self, paragraph, limit=15):
        nlp = spacy.load("en_core_web_sm")
        # Process the paragraph using spaCy
        doc = nlp(paragraph)

        # Extract noun chunks and verbs (potential subjects and actions for questions)
        noun_chunks = [chunk.text for chunk in doc.noun_chunks]
        verbs = [token.text for token in doc if token.pos_ == "VERB"]

        # Generate questions based on extracted noun chunks and verbs
        questions = set()
        count = 0  # Keep track of the number of generated questions

        for chunk in noun_chunks:
            if count == limit:
                break  # Stop generating questions once the limit is reached

            if len(chunk.split()) > 1:
                # Formulate questions for multi-word noun chunks
                questions.add(f"What is {chunk}?")
                questions.add(f"Tell me about {chunk}.")
                questions.add(f"Explain the concept of {chunk}.")
            else:
                # Formulate questions for single-word noun chunks
                questions.add(f"Can you provide information about {chunk}?")

            count += 1

        for verb in verbs:
            if count == limit:
                break

            # Formulate questions related to verbs
            questions.add(f"What does it mean to {verb}?")
            questions.add(f"Describe the process of {verb}.")

            count += 1

        return list(questions)[:limit]  # Convert set to list and limit to 15 questions

    def create_word_document(self, questions):
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        doc_path = os.path.join(desktop_path, "Question_Bank.docx")

        # Create a new Word document
        doc = Document()

        # Add image at the top with a border of 0.8mm
        paragraph = doc.add_paragraph('')
        paragraph.alignment = 1

        # Add the image to the paragraph
        image_path = "C:/Users/Ayaan Deshmukh/Downloads/kalsekar.png"  # Replace with the actual path to your image
        paragraph.add_run().add_picture(image_path, width=Inches(6.3))

        # Add space between image and title
        doc.add_paragraph()

        # Add title to the document
        title = doc.add_paragraph("Question Bank")

        title.runs[0].underline = True

        title.runs[0].bold = True
        title.paragraph_format.alignment = 1  # Center alignment
        title.runs[0].font.size = Pt(18)  # Set font size to 18

        # Add space between title and questions
        doc.add_paragraph()
        sec_pr = doc.sections[0]._sectPr if doc.sections else doc.sections.add()._sectPr

        # Add a border to the paragraph
        pg_borders = OxmlElement('w:pgBorders')
        pg_borders.set(qn('w:offsetFrom'), 'page')

        # Set individual borders
        for border_name in ('top', 'left', 'bottom', 'right'):
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')  # a single line
            border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
            border_el.set(qn('w:space'), '24')  # space
            border_el.set(qn('w:color'), 'auto')  # color
            pg_borders.append(border_el)

        # Append page borders to section properties
        sec_pr.append(pg_borders)

        # Add generated questions to the document
        for i, question in enumerate(questions, start=1):
            doc.add_paragraph(f"{i}. {question}")

        # Save the Word document
        doc.save(doc_path)
        print("Word document created and saved on the desktop.")

    def show_help(self):
        # Display help dialog box
        help_text = "Welcome to Question Maker!\n\n"\
                    "1. Click 'Select PDF' to choose a PDF file.\n"\
                    "2. Wait for a moment while the content is processed.\n"\
                    "3. Explore the generated questions in a Word document on your desktop.\n"\
                    "4. Enjoy learning!\n\n"\
                    "Note: Make sure to have a PDF file ready for processing."
        messagebox.showinfo("How to Use", help_text)

    def explore_modal_paper(self):
        # Action for Modal Answer Paper button
        messagebox.showinfo("Explore Modal Answer Paper", "Explore Official Modal Answer Papers.")

    def explore_question_papers(self):
        # Action for Question Papers button
        messagebox.showinfo("Explore Question Papers", "Explore MSBTE Official Exam Question Papers.")

if __name__ == "__main__":
    nlp = spacy.load("en_core_web_sm")
    root = tk.Tk()
    app = QB(root)
    root.mainloop()
