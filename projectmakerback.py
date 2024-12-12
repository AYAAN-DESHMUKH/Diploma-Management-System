import tkinter as tk
from tkinter import scrolledtext
from docx.shared import Inches, Pt
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from projectmakerui import ProjectInfoUI

def generate_project_document():
    doc = Document()
    
    # Get the section properties element
    sec_pr = doc.sections[0]._sectPr if doc.sections else doc.sections.add()._sectPr

    # Create new page borders element
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')

    # Set individual borders
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')  # a single line
        border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
        border_el.set(qn('w:space'), '24')  # space
        border_el.set(qn('w:color'), 'auto')  # color
        pg_borders.append(border_el)
    
    # Append page borders to section properties
    sec_pr.append(pg_borders)

    # Set the document alignment to center
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            for run in paragraph.runs:
                run.paragraph.alignment = 1

    # Adjust the left margin to center the image
    for section in doc.sections:
        section.left_margin = Inches((8.5 - 7.0) / 2)  # Assuming default page width is 8.5 inches

    # Add a new paragraph with centered alignment
    paragraph = doc.add_paragraph('')
    paragraph.alignment = 1

    # Add the image to the paragraph
    image_path = "C:/Users/Admin/Desktop/kalsekar.png"  # Replace with the actual path to your image
    paragraph.add_run().add_picture(image_path, width=Inches(7.0))

    # Add a new paragraph with project details
    project_title_paragraph = doc.add_paragraph("")
    project_title_paragraph.add_run("A Major-Project Proposal\nOn\n").bold = False
    project_title_paragraph.runs[0].font.italic = True
    project_title_paragraph.paragraph_format.alignment = 1

    # Add the run and set its font size to 22pt
    run_project_title = project_title_paragraph.add_run("\n")
    run_project_title.bold = True
    run_project_title.font.size = Pt(22)

    # Add a new paragraph with team members
    team_members_paragraph = doc.add_paragraph("")
    team_members_paragraph.add_run("By\n").bold = True
    team_members_paragraph.add_run("1. Khan Mohd. Anas\n").bold = False
    team_members_paragraph.add_run("2. Dalvi Talha\n").bold = False
    team_members_paragraph.add_run("3. Khan Mohd. Faiz\n").bold = False
    team_members_paragraph.add_run("4. Deshmukh Ayaan\n").bold = False
    team_members_paragraph.paragraph_format.alignment = 1  # Align center
    team_members_paragraph.paragraph_format.space_after = Pt(18)  # Add spacing

    # Add a new paragraph with guidance
    guidance_paragraph = doc.add_paragraph("")
    guidance_paragraph.add_run("Under Guidance of\n").bold = True
    guidance_paragraph.add_run('"Hafsha Siddique Maam"').bold = False
    guidance_paragraph.paragraph_format.alignment = 1  # Align center
    guidance_paragraph.paragraph_format.space_after = Pt(18)  # Add spacing

    # Add a new paragraph with course details
    course_details_paragraph = doc.add_paragraph("")
    course_details_paragraph.add_run("In\n").bold = True
    course_details_paragraph.add_run("Three Years Diploma Program in Engineering & Technology of\n").bold = False
    course_details_paragraph.add_run("Maharashtra State Board of Technical Education, Mumbai (Autonomous)\n").bold = False
    course_details_paragraph.add_run("ISO 9001:2015\n").bold = False
    course_details_paragraph.paragraph_format.alignment = 1  # Align center
    course_details_paragraph.paragraph_format.space_after = Pt(18)  # Add spacing

    # Set font size for the entire document
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(19)

    """
    
    second page
    
    """

    doc.add_page_break()
    # Get the section properties element

    # Create new page borders element
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')

    # Set individual borders
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')  # a single line
        border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
        border_el.set(qn('w:space'), '24')  # space
        border_el.set(qn('w:color'), 'auto')  # color
        pg_borders.append(border_el)
    # Append page borders to section properties
    sec_pr.append(pg_borders)
    paragraph2 = doc.add_paragraph('')
    paragraph2.alignment = 1
    image_pat = "C:/Users/Admin/Desktop/cert.png"  # Replace with the actual path to your image
    paragraph2.add_run().add_picture(image_pat, width=Inches(7.0))

    """
    third page
    """

    # Create new page borders element
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')

    # Set individual borders
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')  # a single line
        border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
        border_el.set(qn('w:space'), '24')  # space
        border_el.set(qn('w:color'), 'auto')  # color
        pg_borders.append(border_el)
    # Append page borders to section properties
    sec_pr.append(pg_borders)
    paragraph3 = doc.add_paragraph('')
    paragraph3.alignment = 1
    image_pat3 = "C:/Users/Admin/Desktop/annexure1.png"  # Replace with the actual path to your image
    paragraph3.add_run().add_picture(image_pat3, width=Inches(7.0))

    """
    fourth page
    """

    # Create new page borders element
    paragraph3.alignment = 1
    image_pat3 = "C:/Users/Admin/Desktop/ann2.png"  # Replace with the actual path to your image
    paragraph3.add_run().add_picture(image_pat3, width=Inches(7.0))

    """
    fifth page
    """

    # Create new page borders element
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')

    # Set individual borders
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')  # a single line
        border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
        border_el.set(qn('w:space'), '24')  # space
        border_el.set(qn('w:color'), 'auto')  # color
        pg_borders.append(border_el)
    # Append page borders to section properties
    sec_pr.append(pg_borders)
    paragraph3 = doc.add_paragraph('')
    paragraph3.alignment = 1
    image_pat3 = "C:/Users/Admin/Desktop/ann3.png"  # Replace with the actual path to your image
    paragraph3.add_run().add_picture(image_pat3, width=Inches(7.0))

    # new page

    # Add a new section for the next page
    doc.add_section()

    # Get the section properties element for the new section
    new_sec_pr = doc.sections[-1]._sectPr

    # Create new page borders element for the new section
    new_pg_borders = OxmlElement('w:pgBorders')
    new_pg_borders.set(qn('w:offsetFrom'), 'page')

    # Set individual borders for the new section
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')  # a single line
        border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
        border_el.set(qn('w:space'), '24')  # space
        border_el.set(qn('w:color'), 'auto')  # color
        new_pg_borders.append(border_el)

    # Append new page borders to section properties for the new section
    new_sec_pr.append(new_pg_borders)

    # Add a new centered paragraph with the title "Introduction"
    introduction_paragraph = doc.add_paragraph("Introduction")
    introduction_paragraph.paragraph_format.alignment = 1  # Center alignment
    introduction_paragraph.runs[0].font.size = Pt(32)  # Set font size to 32

    # Add a new paragraph below the introduction with font size 12
    text_below_paragraph = doc.add_paragraph("This is additional text below the Introduction.")
    text_below_paragraph.runs[0].font.size = Pt(12)  # Set font size to 12
    paragraph_format = text_below_paragraph.paragraph_format
    paragraph_format.line_spacing = Pt(53)
    doc.add_page_break()

    # new page abstract page
    # Get the section properties element for the new section
    new_sec_pr = doc.sections[-1]._sectPr

    # Create new page borders element for the new section
    new_pg_borders = OxmlElement('w:pgBorders')
    new_pg_borders.set(qn('w:offsetFrom'), 'page')

    # Set individual borders for the new section
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')  # a single line
        border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
        border_el.set(qn('w:space'), '24')  # space
        border_el.set(qn('w:color'), 'auto')  # color
        new_pg_borders.append(border_el)

    # Append new page borders to section properties for the new section
    new_sec_pr.append(new_pg_borders)

    # Add a new centered paragraph with the title "Introduction"
    introduction_paragraph = doc.add_paragraph("Abstract")
    introduction_paragraph.paragraph_format.alignment = 1  # Center alignment
    introduction_paragraph.runs[0].font.size = Pt(32)  # Set font size to 32

    # Add a new paragraph below the introduction with font size 12
    text_below_paragraph = doc.add_paragraph("This is additional text below the Introduction.")
    text_below_paragraph.runs[0].font.size = Pt(12)  # Set font size to 12
    paragraph_format = text_below_paragraph.paragraph_format
    paragraph_format.line_spacing = Pt(53)

    doc.add_page_break()

    new_sec_pr = doc.sections[-1]._sectPr

    # Create new page borders element for the new section
    new_pg_borders = OxmlElement('w:pgBorders')
    new_pg_borders.set(qn('w:offsetFrom'), 'page')

    # Set individual borders for the new section
    for border_name in ('top', 'left', 'bottom', 'right'):

        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')  # a single line
        border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
        border_el.set(qn('w:space'), '24')  # space
        border_el.set(qn('w:color'), 'auto')  # color
        new_pg_borders.append(border_el)

    # Append new page borders to section properties for the new section
    new_sec_pr.append(new_pg_borders)

    # Add a new centered paragraph with the title "Introduction"
    introduction_paragraph = doc.add_paragraph("Reference")
    introduction_paragraph.paragraph_format.alignment = 1  # Center alignment
    introduction_paragraph.runs[0].font.size = Pt(32)  # Set font size to 32

    # Add a new paragraph below the introduction with font size 12
    text_below_paragraph = doc.add_paragraph("This is additional text below the Introduction.")
    text_below_paragraph.runs[0].font.size = Pt(12)  # Set font size to 12
    paragraph_format = text_below_paragraph.paragraph_format
    paragraph_format.line_spacing = Pt(53)
    
    
    doc.add_page_break()

    new_sec_pr = doc.sections[-1]._sectPr

    # Create new page borders element for the new section
    new_pg_borders = OxmlElement('w:pgBorders')
    new_pg_borders.set(qn('w:offsetFrom'), 'page')

    # Set individual borders for the new section
    for border_name in ('top', 'left', 'bottom', 'right'):

        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')  # a single line
        border_el.set(qn('w:sz'), '12')  # border size (4.5 pt)
        border_el.set(qn('w:space'), '24')  # space
        border_el.set(qn('w:color'), 'auto')  # color
        new_pg_borders.append(border_el)

    # Append new page borders to section properties for the new section
    new_sec_pr.append(new_pg_borders)

    # Add a new centered paragraph with the title "Introduction"
    introduction_paragraph = doc.add_paragraph("Conclusion")
    introduction_paragraph.paragraph_format.alignment = 1  # Center alignment
    introduction_paragraph.runs[0].font.size = Pt(32)  # Set font size to 32

    # Add a new paragraph below the introduction with font size 12
    text_below_paragraph = doc.add_paragraph("This is additional text below the Introduction.")
    text_below_paragraph.runs[0].font.size = Pt(12)  # Set font size to 12
    paragraph_format = text_below_paragraph.paragraph_format
    paragraph_format.line_spacing = Pt(53)
    
    
  

